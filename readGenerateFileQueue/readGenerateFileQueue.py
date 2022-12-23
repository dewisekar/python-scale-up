import pyodbc 
import pandas as pd
import os
# from os import path
import platform 
import sys
sys.path.append(os.path.join(os.path.dirname(__file__), '..', 'database'))
sys.path.append(os.path.join(os.path.dirname(__file__), '..', 'logger'))
import database
import myLogger

import pika
import time
import json

from PyPDF2 import PdfFileReader, PdfFileWriter
from PyPDF2.generic import DecodedStreamObject, EncodedStreamObject, NameObject
from docx import Document
import datetime
import convertapi
import shutil 

isNeedConvertToPdf = True
mapMonth = [
    'Januari',
    'Februari',
    'Maret',
    'April',
    'Mei',
    'Juni',
    'Juli',
    'Agustus',
    'September',
    'Oktober',
    'November',
    'Desember'
]

mapDay = {
    'Tue':'Selasa',
    'Wed':'Rabu',
    'Thu':'Kamis',
    'Fri':'Jumat',
    'Sat':'Sabtu',
    'Sun':'Minggu',
    'Mon':'Senin'
}

mapData = {
    "kontrakId":0,
    "kolId":1,
    "managerId":2,
    "fileId":3,
    "managerName":4,
    "managerRoles":5,
    "managerKtp":6,
    "managerAddress":7,
    "kolName":8,
    "kolKtp":9,
    "alamatKol":10,
    "platform":11,
    "username":12,
    "bookingSlot":13,
    "tanggalAwal":14,
    "tanggalAkhir":15,
    "biaya":16,
    "norekKol":17,
    "bankKol":18,
    "dpPercentage":19
}

mapReplacement = {
    "MANAGER_NAME":"managerName",
    "MANAGER_ROLE":"managerRoles",
    "MANAGER_KTP":"managerKtp",
    "MANAGER_ALAMAT":"managerAddress",
    "KOL_NAME":"kolName",
    "KOL_KTP":"kolKtp",
    "KOL_ALAMAT":"alamatKol",
    "PLATFORM":"platform",
    "USERNAME":"username",
    "NOREK":"norekKol",
    "BANK":"bankKol",
}
##############################
def replace_text(content, replacements = dict()):
    lines = content.splitlines()

    result = ""
    in_text = False

    for line in lines:
        if line == "BT":
            in_text = True

        elif line == "ET":
            in_text = False

        elif in_text:
            cmd = line[-2:]
            if cmd.lower() == 'tj':
                replaced_line = line
                for k, v in replacements.items():
                    replaced_line = replaced_line.replace(k, v)
                result += replaced_line + "\n"
            else:
                result += line + "\n"
            continue

        result += line + "\n"

    return result


def process_data(object, replacements):
    data = object.getData()
    decoded_data = data.decode('utf-8')

    replaced_data = replace_text(decoded_data, replacements)

    encoded_data = replaced_data.encode('utf-8')
    if object.decodedSelf is not None:
        object.decodedSelf.setData(encoded_data)
    else:
        object.setData(encoded_data)

def replaceWordsInFile(filename,replacements):
    filename_result = 'NEW_' + filename
    status = False
    try:
        pdf = PdfFileReader(filename)
        writer = PdfFileWriter()
        for page_number in range(0, pdf.getNumPages()):
            page = pdf.getPage(page_number)
            contents = page.getContents()

            if isinstance(contents, DecodedStreamObject) or isinstance(contents, EncodedStreamObject):
                process_data(contents, replacements)
            elif len(contents) > 0:
                for obj in contents:
                    if isinstance(obj, DecodedStreamObject) or isinstance(obj, EncodedStreamObject):
                        streamObj = obj.getObject()
                        process_data(streamObj, replacements)

            # Force content replacement
            page[NameObject("/Contents")] = contents.decodedSelf
            writer.addPage(page)
        # with open(filename_base + ".result.pdf", 'wb') as out_file: filename_result
        with open(filename_result, 'wb') as out_file: 
            writer.write(out_file)
        status = True
    except Exception as e :
        myLogger.logging_error("readGenerateFileQueue","got exception when replaceWordsInFile : ",e)  
    return filename_result,status

def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)

def replaceWordsInDocx(filename,replacements):
    filename_result = replacements['{ID}'] + '_PKS.WMC_'+  replacements['{BULAN}'] +'_' +  replacements['{TAHUN}'] +'.docx'
    # filename_result_pdf = replacements['{ID}'] + '/PKS.WMC/'+  replacements['{BULAN}'] +'/' +  replacements['{TAHUN}'] +'.pdf'
    status = False
    try:
        template_document = Document(filename)

        for variable_key, variable_value in replacements.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        template_document.save(filename_result)
        if(isNeedConvertToPdf):
            filename_result_pdf,status = convert2pdf(os.path.abspath(os.getcwd()) + '/' + filename_result)
            os.remove(os.path.join(os.path.abspath(os.getcwd()), filename_result)) #remove the docx file
            shutil.move(os.path.abspath(os.getcwd()) + '/' + filename_result.replace('docx','pdf'), os.path.abspath(os.getcwd()) + '/docsMou/' + filename_result.replace('docx','pdf')) #copy the pdf file to folder docsMoU
            filename_result = os.path.abspath(os.getcwd()) + '/docsMou/' + filename_result.replace('docx','pdf')#filename_result_pdf
        else:
            status = True
        # status = True
    except Exception as e:
        status = False
        myLogger.logging_error("readGenerateFileQueue","got exception when replaceWordsInDocx : ",e)  
    return filename_result,status

def convert2pdf(fileinput):
    status = False
    try:
        fileoutput = fileinput.replace('docx','pdf')
        convertapi.api_secret='rRJ2YoPWn7IBiFlb'
        result = convertapi.convert('pdf', { 'File': fileinput })
        result.file.save(fileoutput)
        status = True
    except Exception as e:
        myLogger.logging_error("readGenerateFileQueue","convert2pdf: ",e)

    return fileoutput,status
##############################
def create_replacement(detailData):
    replacement = {}
    for key in mapReplacement:
        replacement['{'+key+'}'] = detailData[mapReplacement[key]]

    replacement['{SLOT}'] = str(detailData['bookingSlot'])
    replacement['{ID}'] = str(detailData['fileId']).zfill(3)
    replacement['{BULAN}'] = str(datetime.datetime.now().month).zfill(2)
    replacement['{TAHUN}'] = str(datetime.datetime.now().year).zfill(4)
    replacement['{TANGGAL_AWAL}'] = str(detailData['tanggalAwal'].day) +' ' + mapMonth[detailData['tanggalAwal'].month-1] 
    replacement['{TANGGAL_AKHIR}'] = str(detailData['tanggalAkhir'].day) +' ' + mapMonth[detailData['tanggalAkhir'].month-1] 
    biaya = detailData['biaya']
    dpPercentage = detailData['dpPercentage']
    replacement['{BIAYA}'] = 'Rp ' +  f'{int(biaya):,}'.replace(',','.') 
    replacement['{DP}'] = f'{int((biaya*dpPercentage)/100):,}'.replace(',','.') 
    replacement['{Sisa_DP}'] = f'{int((biaya*(100-dpPercentage))/100):,}'.replace(',','.') 
    replacement['{DP_Percentage}'] = str(detailData['dpPercentage']) + '% '
    replacement['{DATE_NOW}'] = mapDay[datetime.datetime.now().strftime('%a')] +', ' +str(datetime.datetime.now().day) +' ' + mapMonth[datetime.datetime.now().month-1] #+' ' +str(datetime.datetime.now().year)
    
    return replacement

def process_data_queue(data):
    ret = False
    try:
        fileId = data["FileId"]
        updateString = " UPDATE [MARKETING].[dbo].[File Mou] SET [Progress Status] = 1, [Last Update Date] = GETDATE() "+ "WHERE [File Id] = " + str(fileId) + ";"
        database.executeSqlQuery(updateString)

        selectKontrak = " EXEC	[MARKETING].[dbo].[SP_GetDetailKolKontrakToGenerateFileMou] @FileID = " + str(fileId) + ";"
        rows = database.fetchRowsFromDatabase(selectKontrak)
        myLogger.logging_info("readGenerateFileQueue"," rows:",rows)
        detailData = {}
        row = rows[0]
        for key in mapData:
            detailData[key] = row[mapData[key]]
        myLogger.logging_info("readGenerateFileQueue"," detailData:",detailData)
        replacement = create_replacement(detailData)
        myLogger.logging_info("readGenerateFileQueue"," replacement:",replacement)
        
        if detailData['dpPercentage'] == 100:
            filename_result,status = replaceWordsInDocx('template_full.docx',replacement)
        else :
            filename_result,status = replaceWordsInDocx('template.docx',replacement)
        myLogger.logging_info("readGenerateFileQueue"," filename_result:",filename_result)
        myLogger.logging_info("readGenerateFileQueue"," status:",status)
        if status:
            ret = True
            updateString = " UPDATE [MARKETING].[dbo].[File Mou] SET [Progress Status] = 2, [Last Update Date] = GETDATE(), [File Name]='"+ filename_result+ "' WHERE [File Id] = " + str(fileId) + ";"
            database.executeSqlQuery(updateString)
        else:
            updateString = " UPDATE [MARKETING].[dbo].[File Mou] SET [Progress Status] = -1, [Last Update Date] = GETDATE() "+ "WHERE [File Id] = " + str(fileId) + ";"
            database.executeSqlQuery(updateString)

        myLogger.logging_info("readGenerateFileQueue"," status replaceWordsInDocx:",status)
        # row = rows[0]
        # kontrakId = row[0]
        # kolId = row[1]
        # managerId = row[2]
        # managerName = row[4]
        # managerRoles = row[5]
        # managerKtp = row[6]
        # managerAddress = row[7]
        # kolName = row[8]
        # kolKtp = row[9]
        # alamatKol = row[10]
        # platform = row[11]
        # username = row[12]
        # bookingSlot = row[13]
        # tanggalAwal = row[14]
        # tanggalAkhir = row[15]
        # biaya = row[16]
        # dp = biaya/2
        # norekKol = row[17]
        # bankKol = row[18]
    except Exception as e:
        myLogger.logging_error("readGenerateFileQueue","got exception when processing data : ",e)
    return ret     


def consume_queue(queue_name):
    connection = pika.BlockingConnection(pika.ConnectionParameters(host='localhost'))
    channel = connection.channel()

    channel.queue_declare(queue=queue_name, durable=True)

    def callback(ch, method, properties, body):
        myLogger.logging_info("readGenerateFileQueue"," [x] Received %r" % body)
        # print(" [x] Received %r" % body)
        json_object = json.loads(body)
        status = process_data_queue(json_object)
        # print(json_object)
        myLogger.logging_info("readGenerateFileQueue"," json_object:",json_object)
        if status :
            ch.basic_ack(delivery_tag=method.delivery_tag)

    channel.basic_qos(prefetch_count=1)
    channel.basic_consume(queue=queue_name, on_message_callback=callback)

    print(' [*] Waiting for messages. To exit press CTRL+C')
    channel.start_consuming()

if __name__ == '__main__':
    try:
        consume_queue("generate_file_contract")
    except KeyboardInterrupt:
        print('Interrupted')
        try:
            sys.exit(0)
        except SystemExit:
            os._exit(0)